#!/usr/bin/env python3
"""
generate_templates.py — Convert Markdown templates to DOCX and PDF.

Reads Markdown sources from scripts/templates-md/ (and es/ subdirectory),
converts each to PDF (via weasyprint) and DOCX (via pandoc or python-docx),
and writes output to public/templates/ (and public/templates/es/).

Placeholder handling ({{field_name}}):
  - If a JSON data file is provided (--data), placeholders are replaced with values
    from that file. Empty values → visual gap (underline or HTML blank).
  - In PDF mode: empty placeholders become styled HTML gaps (inline-block with border-bottom).
  - In DOCX mode: empty placeholders become literal underscore lines.
  - Without --data (build time): all placeholders become visual gaps, so the PDF/DOCX
    templates show blank spaces ready to be filled manually.
"""
import argparse
import json
import re
import subprocess
import sys
from pathlib import Path

from markdown_it import MarkdownIt

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

ROOT = Path(__file__).resolve().parent.parent
SRC_DIR = ROOT / "scripts" / "templates-md"
OUT_DIR = ROOT / "public" / "templates"

LANGUAGES = [
    ("", "eu"),    # Basque (default) -> public/templates/
    ("es", "es"),  # Spanish        -> public/templates/es/
]

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

md = MarkdownIt("commonmark", {"html": True})
# Enable GFM tables (rule exists but not in commonmark preset)
md.block.ruler.enable("table", True)
md.inline.ruler.enable("table", True)
md.block.ruler.enable("table", True)


def _blank_for_pdf(width_chars: int = 20) -> str:
    """Return an HTML gap that renders as a fillable blank line in weasyprint PDF.

    Uses an inline-block span with a bottom border — looks like a handwritten
    line in the PDF but renders as a visible blank space on screen too.
    """
    width_px = width_chars * 9  # rough: 9px per char at 11pt
    return (
        f'<span style="display:inline-block;min-width:{width_px}px;'
        f'border-bottom:1px solid #666;text-align:center;">&nbsp;</span>'
    )


def _blank_for_docx(width_chars: int = 20) -> str:
    """Return a literal underscore line for DOCX (pandoc converts _ to underscore)."""
    return "_" * width_chars


def replace_placeholders(text: str, data: dict | None, *, for_pdf: bool) -> str:
    """Replace all {{placeholder}} occurrences in text.

    - If data has the key and a non-empty value → replace with the value.
    - If data has the key but empty value → replace with visual blank/gap.
    - If key is missing from data entirely → replace with visual blank/gap.
    - for_pdf=True → HTML gaps; for_pdf=False → literal underscores (DOCX).

    Supports inline placeholders in table cells, paragraphs, headings, etc.
    """
    if data is None:
        data = {}

    # Estimate gap width from the placeholder name length + some padding
    def gap(key: str) -> str:
        # Width roughly proportional to field name, capped between 15-35 chars
        width = max(15, min(35, len(key) + 5))
        return _blank_for_pdf(width) if for_pdf else _blank_for_docx(width)

    pattern = re.compile(r"\{\{(\w+)\}\}")

    def replacer(m: re.Match[str]) -> str:
        key = m.group(1)
        value = data.get(key)
        if value and isinstance(value, str):
            v = value.strip()
            if v:
                return v
        # Empty or missing → visual gap
        return gap(key)

    return pattern.sub(replacer, text)


def md_to_html(md_path: Path, data: dict | None = None) -> str:
    """Read a Markdown file, replace placeholders, and return HTML string."""
    text = md_path.read_text(encoding="utf-8")
    text = replace_placeholders(text, data, for_pdf=True)
    body = md.render(text)

    # Minimal HTML wrapper for weasyprint
    return f"""<!DOCTYPE html>
<html lang="eu">
<head>
<meta charset="utf-8">
<style>
  body {{
    font-family: "DejaVu Sans", sans-serif;
    font-size: 11pt;
    line-height: 1.5;
    max-width: 800px;
    margin: 2rem auto;
    padding: 0 1rem;
    color: #1a1a1a;
  }}
  h1 {{ font-size: 1.6em; border-bottom: 2px solid #2563eb; padding-bottom: .3em; }}
  h2 {{ font-size: 1.3em; margin-top: 1.5em; }}
  h3 {{ font-size: 1.1em; }}
  table {{
    border-collapse: collapse; width: 100%; margin: 1em 0;
    page-break-inside: avoid;
  }}
  thead {{ display: table-header-group; }}
  th, td {{
    border: 1px solid #999; padding: .4em .6em; text-align: left;
    font-size: 10pt;
  }}
  th {{ background: #e5e7eb; font-weight: 600; }}
  code {{ background: #f3f4f6; padding: .2em .4em; border-radius: 3px; font-size: .9em; }}
  pre {{ background: #f3f4f6; padding: 1em; border-radius: 4px; overflow-x: auto; }}
</style>
</head>
<body>
{body}
</body>
</html>"""


def html_to_pdf(html: str, out_path: Path) -> None:
    """Convert HTML string to PDF via weasyprint."""
    from weasyprint import HTML as WHTML
    out_path.parent.mkdir(parents=True, exist_ok=True)
    WHTML(string=html).write_pdf(str(out_path))
    print(f"  PDF  → {out_path.relative_to(ROOT)}")


def html_to_docx(html: str, out_path: Path) -> None:
    """
    Convert HTML to DOCX via pandoc.
    Falls back to a minimal python-docx stub if pandoc is unavailable.
    """
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # Try pandoc first (handles tables, styles, etc. best)
    try:
        proc = subprocess.run(
            [
                "pandoc",
                "--from", "html",
                "--to", "docx",
                "--output", str(out_path),
            ],
            input=html,
            text=True,
            capture_output=True,
            timeout=30,
        )
        if proc.returncode == 0 and out_path.exists():
            print(f"  DOCX → {out_path.relative_to(ROOT)}")
            return
        else:
            print(f"  ⚠ pandoc failed for {out_path.name}: {proc.stderr[:200]}", file=sys.stderr)
    except FileNotFoundError:
        print("  ⚠ pandoc not found, using python-docx fallback.", file=sys.stderr)

    # Fallback: python-docx (basic, no table support)
    from docx import Document
    doc = Document()
    doc.add_heading(out_path.stem.replace("_", " ").title(), 0)
    doc.add_paragraph(
        "Txantiloi hau automatikoki sortu da. Erabili Markdown jatorrizkoa formato "
        "aberatsago baterako edo ireki DOCX hau eta editatu zure beharretara."
    )
    doc.save(str(out_path))
    print(f"  DOCX → {out_path.relative_to(ROOT)} (python-docx fallback)")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="Generate DOCX/PDF templates from Markdown")
    parser.add_argument("--lang", choices=["eu", "es"], help="Only process one language")
    parser.add_argument(
        "--data",
        type=Path,
        help="JSON file with user data for placeholder replacement (e.g., from localStorage export)",
    )
    args = parser.parse_args()

    # Load optional user data
    user_data: dict | None = None
    if args.data:
        data_path = Path(args.data)
        if data_path.exists():
            user_data = json.loads(data_path.read_text(encoding="utf-8"))
            print(f"📋 Loaded user data from {data_path}")
        else:
            print(f"⚠ Data file not found: {data_path}", file=sys.stderr)

    processed = 0
    failed = 0

    for subdir, lang in LANGUAGES:
        if args.lang and lang != args.lang:
            continue

        src = SRC_DIR / subdir if subdir else SRC_DIR
        dst = OUT_DIR / subdir if subdir else OUT_DIR

        if not src.exists():
            print(f"⚠ {src.relative_to(ROOT)} does not exist, skipping {lang}")
            continue

        for md_file in sorted(src.glob("*.md")):
            name = md_file.stem
            print(f"▶ {name} ({lang})")

            try:
                # PDF: placeholders → HTML gaps
                html = md_to_html(md_file, user_data)
                html_to_pdf(html, dst / f"{name}.pdf")
                # DOCX: placeholders → literal underscore lines
                md_text = md_file.read_text(encoding="utf-8")
                md_text = replace_placeholders(md_text, user_data, for_pdf=False)
                html_for_docx = md.render(md_text)
                html_to_docx(html_for_docx, dst / f"{name}.docx")
                processed += 1
            except Exception as exc:
                print(f"  ✗ FAILED: {exc}", file=sys.stderr)
                import traceback
                traceback.print_exc()
                failed += 1

    print(f"\nDone. {processed} file(s) processed, {failed} failed.")
    sys.exit(1 if failed else 0)


if __name__ == "__main__":
    main()
